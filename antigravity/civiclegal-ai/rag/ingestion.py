"""
CivicLegal AI — Document Ingestion Pipeline
Loads, chunks, and indexes legal documents into ChromaDB + FAISS.
"""
import os
import logging
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)

CHUNK_SIZE = 500
CHUNK_OVERLAP = 75


def load_pdf(file_path: str) -> List[dict]:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                pages.append({"text": text, "page": i + 1, "source": Path(file_path).name})
        doc.close()
        return pages
    except Exception as e:
        logger.error(f"PDF load error {file_path}: {e}")
        return []


def load_docx(file_path: str) -> List[dict]:
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [{"text": text, "page": 1, "source": Path(file_path).name}]
    except Exception as e:
        logger.error(f"DOCX load error {file_path}: {e}")
        return []


def load_text(file_path: str) -> List[dict]:
    try:
        text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
        return [{"text": text, "page": 1, "source": Path(file_path).name}]
    except Exception as e:
        logger.error(f"Text load error {file_path}: {e}")
        return []


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk)
        start = end - overlap
    return chunks


def ingest_directory(directory: str, vectorstore=None) -> int:
    directory = Path(directory)
    if not directory.exists():
        logger.warning(f"Ingestion directory not found: {directory}")
        return 0

    total_chunks = 0
    for file_path in directory.rglob("*"):
        if file_path.suffix.lower() == ".pdf":
            pages = load_pdf(str(file_path))
        elif file_path.suffix.lower() == ".docx":
            pages = load_docx(str(file_path))
        elif file_path.suffix.lower() in {".txt", ".md"}:
            pages = load_text(str(file_path))
        else:
            continue

        for page_data in pages:
            chunks = chunk_text(page_data["text"])
            for chunk in chunks:
                if vectorstore:
                    try:
                        from langchain.schema import Document
                        vectorstore.add_documents([
                            Document(
                                page_content=chunk,
                                metadata={
                                    "source": page_data["source"],
                                    "page": page_data["page"],
                                    "title": page_data["source"],
                                },
                            )
                        ])
                    except Exception as e:
                        logger.error(f"Vectorstore add error: {e}")
                total_chunks += 1

    logger.info(f"Ingested {total_chunks} chunks from {directory}")
    return total_chunks


def build_bm25_index(texts: List[str]):
    try:
        from rank_bm25 import BM25Okapi
        tokenized = [t.lower().split() for t in texts]
        return BM25Okapi(tokenized)
    except ImportError:
        logger.warning("rank_bm25 not installed — BM25 index skipped")
        return None
